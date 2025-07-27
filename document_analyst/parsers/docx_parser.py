"""
DOCX Parser Module
Handles parsing of Microsoft Word documents.
"""

from docx import Document
from typing import Dict, Any

class DOCXParser:
    """Parser for DOCX documents."""
    
    def __init__(self):
        """Initialize the DOCX parser."""
        pass
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a DOCX file and extract content.
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            dict: Parsed content with metadata
        """
        try:
            doc = Document(file_path)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'creator': core_props.author or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'num_paragraphs': len(doc.paragraphs)
            }
            
            # Extract text content
            full_text = ""
            paragraphs = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if para_text:  # Only include non-empty paragraphs
                    full_text += para_text + "\n"
                    paragraphs.append({
                        'paragraph_number': i + 1,
                        'content': para_text,
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
            
            # Extract tables if any
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                tables.append({
                    'table_number': i + 1,
                    'data': table_data
                })
                
                # Add table content to full text
                for row in table_data:
                    full_text += " | ".join(row) + "\n"
            
            return {
                'content': full_text,
                'metadata': metadata,
                'paragraphs': paragraphs,
                'tables': tables,
                'file_type': 'docx'
            }
            
        except Exception as e:
            raise Exception(f"Error parsing DOCX file {file_path}: {str(e)}")
    
    def is_supported(self, file_path: str) -> bool:
        """Check if the file is a supported DOCX."""
        return file_path.lower().endswith(('.docx', '.doc'))
